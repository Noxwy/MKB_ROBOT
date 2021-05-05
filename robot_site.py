# -*- coding: utf-8 -*-

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time, requests
import os
import threading
from docx import Document
from docx.shared import Inches
from wordautosave import word_automate


url1 = 'https://alfabank.ru/' #Альфа банк
url3 = 'https://mkb.ru/' #МКБ банк
url4 = 'https://www.vtb.ru/' #ВТБ банк

driver = 'chromedriver.exe'
path = 'path to dir'

def alfa_bank():
    #Проверка на существоание директории Альфа
    if os.path.exists(r'C:\Users\...\ALFA')==False:
       os.makedirs(path+'\\ALFA')
       
    browser = webdriver.Chrome(driver)
    browser.set_window_position(0, 100)
    browser.set_window_size(1920, 150)
    browser.get(url1)
    
    #Нажатие конопки Карты    
    url_cards = browser.find_element_by_xpath('//*[@id="alfa"]/div/div[1]/div/div/div/div[2]/div[2]/span[1]/span/a').get_attribute("href")
    browser.get(url_cards)
    
    #Нажатие конопки Карты-> Дебетовые
    url_debet_cards = browser.find_element_by_xpath('//*[@id="alfacard-benefit"]/div/div/div/div[4]/h2/span/a').get_attribute("href")
    browser.get(url_debet_cards)
    
    #Получение инофрмации по картам
    text_card = browser.find_element_by_id('all-cards')
    
    #Сохранение в txt
    with open(path+"\\ALFA\\Output_alfa.txt","w") as file:
            file.write(text_card.text.replace('₽','руб')) 

    ##Сохранение изображения карты        
    list_element_img = browser.find_elements_by_class_name('c2g8fdV')
    
    count = 1
    for item in list_element_img:
        #print(item.get_attribute("src"))
        
        #Сохранение в JPG
        try:
            response = requests.get(item.get_attribute("src"))
            with open(path+"\\ALFA\\img_alfa{i}.jpg".format(i = str(count)),"wb") as file_img:
                file_img.write(response.content)
        except:
            response = requests.get(item.get_attribute("data-src"))
            with open(path+"\\ALFA\\img_alfa{i}.jpg".format(i = str(count)),"wb") as file_img:
                file_img.write(response.content)
        count+=1
        
    browser.close()
    
    
def mkb_bank():
    #Проверка на существоание директории МКБ
    if os.path.exists(r'C:\Users\...\MKB')==False:
        os.makedirs(path+'\\MKB')
        
    browser = webdriver.Chrome(driver)
    browser.set_window_position(0, 250)
    browser.set_window_size(1920, 150)
    browser.get(url3)
    
    url_cards = browser.find_element_by_xpath('/html/body/header/div[3]/div[2]/div[2]/ul/a[3]').get_attribute("href")
    browser.get(url_cards)
    
    #Получение Текста через перебор массива от 1 до 3
    for a in range(1,3):
        res = browser.find_element_by_xpath('//*[@id="page"]/div[2]/div/div[2]/div/div[1]/div[2]/div/div[1]/div/div[{i}]'.format(i = str(a)))
        with open(path+"\\MKB\\Output_mkb_{i}.txt".format(i=str(a)), "w") as text_file:
            text_file.write(str(res.text.replace('₶','руб')))
    
    #Сохранение изображения карты 
    for a in range(1,3):
        img_url = browser.find_element_by_xpath('//*[@id="page"]/div[2]/div/div[2]/div/div[1]/div[2]/div/div[1]/div/div[1]/div/div[1]/div/div/a/img[{i}]'.format(i=str(a))).get_attribute("src")
        response = requests.get(img_url)
        
        #Сохранение в JPG
        with open(path+"\\MKB\\img_mkb{i}.jpg".format(i = str(a)),"wb") as file:
                    file.write(response.content)
    #Закрытие браузера
    browser.close()
    

def vtb_bank():
    #Проверка на существоание директории ВТБ
    if os.path.exists(r'C:\Users\...\VTB')==False:
        os.makedirs(path+'\\VTB')
        
    browser = webdriver.Chrome(driver) 
    browser.set_window_position(0, 400)
    browser.set_window_size(1920, 500)
    #/html/body/div[4]/div[2]/div/div[2]/div/div[2]/div/div/span/div/div/div/div[2]/a[2]/span
    browser.get(url4)

    #Нажатие конопки Карты   
    browser.find_element_by_xpath('//*[@id="header"]/header/div[1]/div/div/div[2]/a[2]/span').click()

    #Нажатие конопки Карты-> Дебетовые    
    browser.find_element_by_xpath('//*[@id="header"]/header/div[2]/div/div/div[1]/div/div[2]/div[1]/a[2]/span').click()

    #Получение инофрмации по картам
    list_text = browser.find_element_by_xpath('/html/body/main/div/section[4]/div/div').text

    #Сохранение в txt
    with open(path+"\\VTB\\Output_vtb.txt","w") as file:
            file.write(list_text.replace('₽','руб'))
    #Сохранение изображения карты        
    for a in range(1,6):
        img = browser.find_element_by_xpath('/html/body/main/div/section[4]/div/div/div[{i}]/div[1]/img'.format(i = str(a))).get_attribute("src")
        response = requests.get(img)
        
        #Сохранение в JPG
        with open(path+"\\VTB\\img_vtb{i}.jpg".format(i = str(a)),"wb") as file:
            file.write(response.content)
            
    browser.close()
    #list_img = browser.find_elements_by_css_selector('.common-content__col_3.common-content__col_3-lg-screen-wider.common-content__col_image-block')

def writetoDocVTB():
    
    
    document = Document()
    
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(path+'/VTB/img_vtb1.jpg', width=Inches(3.00))
    document.add_paragraph(
    'Дебетовая Мультикарта ВТБ', style='Intense Quote'
)
    document.add_paragraph(
    'Бесплатное обслуживание карты', style='List Bullet'
)
    document.add_paragraph(
    'Кешбэк до 1,5%', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатные переводы в другие банки', style='List Bullet'
)
    r.add_picture(path+'/VTB/img_vtb2.jpg', width=Inches(3.00))
    document.add_paragraph(
    'Цифровая Мультикарта ВТБ', style='Intense Quote'
)
    document.add_paragraph(
    'Моментальный выпуск в ВТБ Онлайн', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатное обслуживание карты', style='List Bullet'
)
    document.add_paragraph(
    'Кешбэк до 1,5%', style='List Bullet'
) 
    # r.add_text('Good Morning every body,This is my')
    #r.add_picture(path+'/img_vtb1.jpg', width=Inches(3.00))
    #r.add_text(' \ndo you like it?')
    
    document.save('VTB.docx')
    
def writetoDocMKB():
    document = Document()
    
    p = document.add_paragraph()
    r = p.add_run()
    
    r.add_picture(path+'/MKB/img_mkb1.jpg', width=Inches(3.00))
    
    document.add_paragraph(
    'Москарта', style='Intense Quote'
)
    document.add_paragraph(
    'Бесплатные переводы до 25 тыс. руб', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатное снятие наличных', style='List Bullet'
)
    document.add_paragraph(
    '5% от покупок в 2 категориях', style='List Bullet'
)
    document.add_paragraph(
    'Сервис мультивалютности', style='List Bullet'
)

    r.add_picture(path+'/MKB/img_mkb2.jpg', width=Inches(3.00))
    
    document.add_paragraph(
    'Москарта Black', style='Intense Quote'
)
    document.add_paragraph(
    'Бесплатные переводы до 50 тыс. руб', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатное снятие наличных', style='List Bullet'
)
    document.add_paragraph(
    '7% от покупок в 3 категориях', style='List Bullet'
)
    document.add_paragraph(
    'Сервис мультивалютности', style='List Bullet'
)
    document.add_paragraph(
    'Страхование путешественников', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатное посещение бизнес-залов', style='List Bullet'
) 
    document.save('MKB.docx')
    
def writetoDocALFA():
    document = Document()
    
    p = document.add_paragraph()
    r = p.add_run()
    
    r.add_picture(path+'/ALFA/img_alfa1.jpg', width=Inches(3.00))
    document.add_paragraph(
    'Дебетовая Альфа-Карта', style='Intense Quote'
)
    document.add_paragraph(
    'Бесплатная всегда', style='List Bullet'
)
    document.add_paragraph(
    'До 2% кэшбэк на покупки', style='List Bullet'
)
    document.add_paragraph(
    'До 5% годовых на остаток по карте', style='List Bullet'
)
    document.add_paragraph(
    ' Бесплатно выпуск и обслуживание', style='List Bullet'
)

    r.add_picture(path+'/ALFA/img_alfa2.jpg', width=Inches(3.00))
    document.add_paragraph(
    'Дебетовая Альфа-Карта Premium', style='Intense Quote'
)
    document.add_paragraph(
    'Особые привилегии премиального обслуживания', style='List Bullet'
)
    document.add_paragraph(
    'До 3% кэшбэк на покупки', style='List Bullet'
)
    document.add_paragraph(
    'До 6% годовых на остаток по карте', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатно снятие наличных', style='List Bullet'
)

    r.add_picture(path+'/ALFA/img_alfa3.jpg', width=Inches(3.00))
    document.add_paragraph(
    'Дебетовая карта Alfa Travel', style='Intense Quote'
)
    document.add_paragraph(
    'Самая выгодная банковская карта для путешествий', style='List Bullet'
)
    document.add_paragraph(
    'До 9% милями за покупки на travel.alfabank.ru', style='List Bullet'
)
    document.add_paragraph(
    'До 3% милями за покупки', style='List Bullet'
)
    document.add_paragraph(
    'Бесплатно выпуск и обслуживание', style='List Bullet'
)

    document.save('ALFA.docx')
#Создание потока для ВТБ
x_mkb = threading.Thread(target = vtb_bank)
#Создание потока для МКБ
x_alfa = threading.Thread(target = alfa_bank)
#Создание потока для Альфа
x_vtb = threading.Thread(target = mkb_bank)

#Запуск потоков
x_vtb.start()
x_mkb.start()
x_alfa.start()

#Прекращение работы потоков
x_vtb.join()
x_mkb.join()
x_alfa.join()
#Создание отчета WORD для ВТБ
writetoDocVTB()
#Создание отчета WORD для МКБ
writetoDocMKB()
#Создание отчета WORD для Альфа
writetoDocALFA()

#Переименуем файлы с помощью pywinauto
word_automate(['C:\\Users\...\ALFA.docx',
               'C:\\Users\...\MKB.docx',
               'C:\\Users\...\VTB.docx'])
